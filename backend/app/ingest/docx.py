import uuid
from docx import Document
from io import BytesIO

from app.ingest.chunker import chunk_text
from app.rag.vectorstore import VectorStore


async def ingest_docx(
    content: bytes,
    filename: str,
    vector_store: VectorStore,
    user_id: str = "default"
) -> tuple[str, int]:
    """
    Ingest a .docx file.

    Args:
        content: File content as bytes
        filename: Original filename
        vector_store: VectorStore instance
        user_id: User ID for data isolation

    Returns:
        tuple: (source_id, number_of_chunks)
    """
    source_id = str(uuid.uuid4())

    # Parse DOCX
    doc = Document(BytesIO(content))
    
    # Extract text from paragraphs
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    
    full_text = "\n\n".join(text_parts)
    
    if not full_text or len(full_text) < 50:
        raise ValueError("Could not extract meaningful content from DOCX file")

    # Chunk the text
    chunks = chunk_text(full_text)

    for chunk in chunks:
        chunk["metadata"]["source_id"] = source_id
        chunk["metadata"]["source_name"] = filename
        chunk["metadata"]["source_type"] = "docx"
        chunk["metadata"]["file_path"] = filename

    # Add to vector store
    await vector_store.add_documents(chunks, source_id, filename, "docx", user_id)

    return source_id, len(chunks)
